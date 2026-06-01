# -*- coding: utf-8 -*-
"""
WirkVektor — Geschäftsdokumente (.docx).
Erzeugt 16 Word-Vorlagen in drei Kategorien, alle mit gemeinsamem Briefbogen
(Kopf/Fuß) und Markenstil gemäß DESIGN.md. Platzhalter in [eckigen Klammern].
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import brand as B

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.normpath(os.path.join(HERE, "..", "Vorlagen"))

# Farben als RGBColor
NAVY  = RGBColor.from_string(B.NAVY_DEEP)
TEAL  = RGBColor.from_string(B.VECTOR_TEAL)
SLATE = RGBColor.from_string(B.SLATE_MID)
LIGHT = RGBColor.from_string(B.SLATE_LIGHT)
BODY  = RGBColor.from_string("334155")   # Slate-700, Fließtext
WHITE = RGBColor.from_string(B.WHITE)
BORDER = B.SAFETY_BORDER
TEALHEX = B.VECTOR_TEAL
NAVYHEX = B.NAVY_DEEP

DISPLAY = B.FONT_DISPLAY
BODYFONT = B.FONT_BODY
CO = B.COMPANY

# ---------------------------------------------------------------------------
# Low-level Helfer
# ---------------------------------------------------------------------------
def _set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def _set_cell_margins(cell, top=60, bottom=60, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('bottom', bottom), ('start', left), ('end', right)):
        el = OxmlElement('w:' + tag)
        el.set(qn('w:w'), str(val)); el.set(qn('w:type'), 'dxa')
        m.append(el)
    tcPr.append(m)

def _no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + edge); el.set(qn('w:val'), 'none')
        borders.append(el)
    tblPr.append(borders)

def _row_bottom_border(cell, color=BORDER, sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    el = OxmlElement('w:bottom')
    el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(sz))
    el.set(qn('w:space'), '0'); el.set(qn('w:color'), color)
    borders.append(el); tcPr.append(borders)

def _p_border(p, color=TEALHEX, sz=12, edge='bottom', space=4):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    el = OxmlElement('w:' + edge)
    el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(sz))
    el.set(qn('w:space'), str(space)); el.set(qn('w:color'), color)
    pbdr.append(el); pPr.append(pbdr)

def run(p, txt, size=10.5, color=BODY, bold=False, italic=False, font=BODYFONT,
        caps=False, spacing=None):
    r = p.add_run(txt)
    r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    r.font.name = font; r.font.color.rgb = color
    # East-Asian-Name sicherstellen, damit der Font greift
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:cs'), font)
    if caps:
        r.font.all_caps = True
    if spacing is not None:
        rPr_sp = OxmlElement('w:spacing'); rPr_sp.set(qn('w:val'), str(spacing))
        rPr.append(rPr_sp)
    return r

def para(doc_or_cell, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, space_before=0,
         line=1.15):
    p = doc_or_cell.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after); pf.space_before = Pt(space_before)
    pf.line_spacing = line
    return p

def eyebrow(doc, label, color=TEAL):
    p = para(doc, space_after=2)
    run(p, label.upper(), size=8.5, color=color, bold=True, caps=True, spacing=20)
    return p

def h1(doc, txt, color=NAVY):
    p = para(doc, space_after=4, space_before=2)
    run(p, txt, size=22, color=color, bold=True, font=DISPLAY)
    return p

def h2(doc, txt, color=NAVY):
    p = para(doc, space_after=3, space_before=10)
    run(p, txt, size=13, color=color, bold=True, font=DISPLAY)
    return p

def body(doc, txt, size=10.5, color=BODY, space_after=6, bold=False):
    p = para(doc, space_after=space_after)
    run(p, txt, size=size, color=color, bold=bold)
    return p

def small(doc, txt, color=LIGHT, size=8.5):
    p = para(doc, space_after=2)
    run(p, txt, size=size, color=color)
    return p

def placeholder_note(doc):
    p = para(doc, space_before=10, space_after=0)
    run(p, "Hinweis: Felder in [eckigen Klammern] vor dem Versand ersetzen.",
        size=8.5, color=LIGHT, italic=True)

# ---------------------------------------------------------------------------
# Briefbogen (Kopf-/Fußzeile) — auf jedem Dokument
# ---------------------------------------------------------------------------
def _build_header(section):
    header = section.header
    header.is_linked_to_previous = False
    # bestehenden leeren Absatz nutzen, Tabelle einfügen
    tbl = header.add_table(rows=1, cols=2, width=Cm(17))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_table_borders(tbl)
    left, right = tbl.rows[0].cells
    left.width = Cm(9); right.width = Cm(8)
    # links: Wortmarke + Claim
    lp = left.paragraphs[0]; lp.paragraph_format.space_after = Pt(0)
    run(lp, "WirkVektor", size=17, color=NAVY, bold=True, font=DISPLAY)
    lp2 = left.add_paragraph(); lp2.paragraph_format.space_after = Pt(0)
    run(lp2, CO["claim"], size=8, color=TEAL, bold=True)
    # rechts: Kontakt
    rp = right.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(0)
    run(rp, CO["email"], size=8.5, color=SLATE)
    rp2 = right.add_paragraph(); rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp2.paragraph_format.space_after = Pt(0)
    run(rp2, CO["web"], size=8.5, color=SLATE)
    # Teal-Trennlinie unter dem Kopf
    rule = header.add_paragraph(); rule.paragraph_format.space_before = Pt(2)
    rule.paragraph_format.space_after = Pt(0)
    _p_border(rule, color=TEALHEX, sz=12)

def _build_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    rule = footer.paragraphs[0]
    rule.paragraph_format.space_after = Pt(3)
    _p_border(rule, color=BORDER, sz=6, edge='top', space=4)
    line1 = (f"{CO['name']} · {CO['owner']}, {CO['role']} · "
             f"{CO['street']}, {CO['zip_city']}")
    line2 = (f"{CO['email']} · {CO['web']} · {CO['phone']} · "
             f"{CO['ustid']} · {CO['bank']} {CO['iban']}")
    for ln in (line1, line2):
        p = footer.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run(p, ln, size=7.5, color=LIGHT)

def new_doc(letterhead=True):
    doc = Document()
    # Stile
    normal = doc.styles['Normal']
    normal.font.name = BODYFONT; normal.font.size = Pt(10.5)
    normal.font.color.rgb = BODY
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)
    # Ränder
    for sec in doc.sections:
        sec.left_margin = Cm(2.4); sec.right_margin = Cm(2.4)
        sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(1.8)
        sec.header_distance = Cm(1.1); sec.footer_distance = Cm(0.9)
        if letterhead:
            _build_header(sec); _build_footer(sec)
    return doc

def address_block(doc, anrede="[Firma / Anrede]"):
    """Anschriftenfeld + Datum rechts."""
    tbl = doc.add_table(rows=1, cols=2)
    _no_table_borders(tbl)
    tbl.columns[0].width = Cm(10); tbl.columns[1].width = Cm(7)
    a, d = tbl.rows[0].cells
    a.width = Cm(10); d.width = Cm(7)
    ap = a.paragraphs[0]; run(ap, anrede, size=10.5, color=BODY)
    for ln in ["[Ansprechpartner:in]", "[Straße Hausnummer]", "[PLZ Ort]"]:
        p = a.add_paragraph(); p.paragraph_format.space_after = Pt(0)
        run(p, ln, size=10.5, color=BODY)
    dp = d.paragraphs[0]; dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(dp, "[Ort], [TT.MM.JJJJ]", size=10.5, color=SLATE)
    para(doc, space_after=2)

def info_table(doc, rows, label_w=4.5, value_w=12.5, header=None):
    """Zweispaltige Label/Wert-Tabelle (randlos, dünne Trennlinien)."""
    tbl = doc.add_table(rows=0, cols=2)
    _no_table_borders(tbl)
    for i, (k, v) in enumerate(rows):
        r = tbl.add_row().cells
        r[0].width = Cm(label_w); r[1].width = Cm(value_w)
        _set_cell_margins(r[0]); _set_cell_margins(r[1])
        kp = r[0].paragraphs[0]; kp.paragraph_format.space_after = Pt(0)
        run(kp, k, size=9, color=SLATE, bold=True)
        vp = r[1].paragraphs[0]; vp.paragraph_format.space_after = Pt(0)
        run(vp, v, size=10.5, color=BODY)
        _row_bottom_border(r[0]); _row_bottom_border(r[1])
    para(doc, space_after=2)
    return tbl

def items_table(doc, headers, rows, widths, totals=None):
    """Positionstabelle mit Navy-Kopf, Zebra-Zeilen, optionalen Summenzeilen."""
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_table_borders(tbl)
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].width = Cm(widths[i]); _set_cell_bg(hdr[i], NAVYHEX)
        _set_cell_margins(hdr[i])
        p = hdr[i].paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        if i >= len(headers) - 2 and len(headers) > 2:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run(p, h, size=9, color=WHITE, bold=True)
    for ri, row in enumerate(rows):
        cells = tbl.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].width = Cm(widths[ci])
            _set_cell_margins(cells[ci])
            _set_cell_bg(cells[ci], "FFFFFF" if ri % 2 == 0 else B.OFF_WHITE)
            cells[ci].vertical_anchor = WD_ALIGN_VERTICAL.CENTER
            p = cells[ci].paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            if ci >= len(headers) - 2 and len(headers) > 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run(p, val, size=10, color=BODY if ci == 1 or ci == 0 else BODY)
            _row_bottom_border(cells[ci])
    if totals:
        for ti, (label, value, strong) in enumerate(totals):
            cells = tbl.add_row().cells
            # label über vorletzte Spalten mergen
            n = len(headers)
            merge_to = cells[0].merge(cells[n - 2])
            _set_cell_margins(merge_to)
            lp = merge_to.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            lp.paragraph_format.space_after = Pt(0)
            run(lp, label, size=10, color=NAVY if strong else SLATE, bold=strong)
            vc = cells[n - 1]; _set_cell_margins(vc)
            if strong:
                _set_cell_bg(vc, B.TEAL_TINT)
            vp = vc.paragraphs[0]; vp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            vp.paragraph_format.space_after = Pt(0)
            run(vp, value, size=10.5 if strong else 10,
                color=TEAL if strong else BODY, bold=strong)
    para(doc, space_after=2)
    return tbl

def signature(doc, name=None, role=None):
    para(doc, space_before=10, space_after=18)
    p = para(doc, space_after=0)
    run(p, "Mit freundlichen Grüßen", size=10.5, color=BODY)
    para(doc, space_after=0)
    sig = para(doc, space_after=0)
    run(sig, name or CO["owner"], size=10.5, color=NAVY, bold=True)
    rp = para(doc, space_after=0)
    run(rp, role or CO["role"] + ", " + CO["name"], size=9.5, color=SLATE)

def status_pill(cell, text, color):
    _set_cell_bg(cell, color)
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run(p, text, size=9, color=WHITE, bold=True)

def save(doc, category, filename):
    d = os.path.join(ROOT, category)
    os.makedirs(d, exist_ok=True)
    path = os.path.join(d, filename)
    doc.save(path)
    print("  ✓", category + "/" + filename)

# ===========================================================================
# KATEGORIE 1 — GESCHÄFTSAUSSTATTUNG
# ===========================================================================
GA = "Geschäftsausstattung"

def doc_briefbogen():
    doc = new_doc()
    address_block(doc)
    eyebrow(doc, "Betreff")
    p = para(doc, space_after=10); run(p, "[Betreff des Schreibens]", size=12, color=NAVY, bold=True, font=DISPLAY)
    body(doc, "Sehr geehrte Damen und Herren,")
    body(doc, "[Einleitender Absatz. Sachlich, klar und auf den Punkt — ohne "
              "Übertreibung. Der Briefbogen dient als Grundlage für alle "
              "offiziellen Schreiben von WirkVektor.]")
    body(doc, "[Hauptteil mit dem eigentlichen Anliegen. Bei Bedarf in mehrere "
              "Absätze gegliedert, damit der Text gut lesbar bleibt.]")
    body(doc, "[Abschließender Absatz mit klarer Handlungsaufforderung oder "
              "Ausblick auf die nächsten Schritte.]")
    signature(doc)
    placeholder_note(doc)
    save(doc, GA, "Briefbogen.docx")

def doc_angebot():
    doc = new_doc()
    eyebrow(doc, "Angebot")
    h1(doc, "Angebot Nr. [A-2026-001]")
    address_block(doc)
    info_table(doc, [
        ("Angebotsnummer", "[A-2026-001]"),
        ("Datum", "[TT.MM.JJJJ]"),
        ("Gültig bis", "[TT.MM.JJJJ]"),
        ("Ansprechpartner", CO["owner"]),
    ])
    body(doc, "Sehr geehrte:r [Name],", space_after=6)
    body(doc, "vielen Dank für Ihr Interesse. Gerne unterbreiten wir Ihnen das "
              "folgende Angebot für die strukturierte Einführung von KI in Ihrem "
              "Unternehmen.")
    h2(doc, "Leistungen")
    items_table(doc,
        ["Pos.", "Leistung", "Menge", "Einzelpreis", "Gesamt"],
        [["1", "KI-Readiness-Check (Analyse & Bericht)", "1", "[X.XXX,00 €]", "[X.XXX,00 €]"],
         ["2", "Use-Case-Sprint (Workshop, 2 Tage)", "1", "[X.XXX,00 €]", "[X.XXX,00 €]"],
         ["3", "Governance-Starterpaket", "1", "[X.XXX,00 €]", "[X.XXX,00 €]"]],
        widths=[1.2, 8.3, 1.7, 2.8, 3.0],
        totals=[("Zwischensumme (netto)", "[XX.XXX,00 €]", False),
                ("zzgl. 19 % USt.", "[X.XXX,00 €]", False),
                ("Gesamtbetrag", "[XX.XXX,00 €]", True)])
    h2(doc, "Konditionen")
    body(doc, "Zahlungsziel 14 Tage nach Rechnungsstellung ohne Abzug. Das "
              "Angebot ist freibleibend und bis zum oben genannten Datum gültig. "
              "Es gelten unsere Allgemeinen Geschäftsbedingungen.", size=9.5, color=SLATE)
    signature(doc)
    placeholder_note(doc)
    save(doc, GA, "Angebot.docx")

def doc_rechnung():
    doc = new_doc()
    eyebrow(doc, "Rechnung")
    h1(doc, "Rechnung Nr. [R-2026-001]")
    address_block(doc)
    info_table(doc, [
        ("Rechnungsnummer", "[R-2026-001]"),
        ("Rechnungsdatum", "[TT.MM.JJJJ]"),
        ("Leistungszeitraum", "[TT.MM.–TT.MM.JJJJ]"),
        ("Kundennummer", "[K-XXX]"),
    ])
    body(doc, "Sehr geehrte:r [Name],", space_after=6)
    body(doc, "für die erbrachten Leistungen erlauben wir uns, wie folgt zu "
              "berechnen:")
    items_table(doc,
        ["Pos.", "Leistung", "Menge", "Einzelpreis", "Gesamt"],
        [["1", "KI-Readiness-Check", "1", "[X.XXX,00 €]", "[X.XXX,00 €]"],
         ["2", "Beratungstage à [Tagessatz]", "[X]", "[X.XXX,00 €]", "[X.XXX,00 €]"]],
        widths=[1.2, 8.3, 1.7, 2.8, 3.0],
        totals=[("Nettobetrag", "[XX.XXX,00 €]", False),
                ("19 % USt.", "[X.XXX,00 €]", False),
                ("Rechnungsbetrag", "[XX.XXX,00 €]", True)])
    h2(doc, "Zahlung")
    body(doc, "Bitte überweisen Sie den Rechnungsbetrag innerhalb von 14 Tagen "
              "ohne Abzug auf das unten genannte Konto unter Angabe der "
              "Rechnungsnummer.", size=9.5, color=SLATE)
    info_table(doc, [
        ("Bank", CO["bank"]),
        ("IBAN", CO["iban"]),
        ("BIC", CO["bic"]),
        ("Verwendungszweck", "[R-2026-001]"),
    ])
    small(doc, "Hinweis bei Bedarf: Gemäß § 19 UStG wird keine Umsatzsteuer "
               "berechnet. (Nur falls Kleinunternehmerregelung — sonst entfernen.)")
    placeholder_note(doc)
    save(doc, GA, "Rechnung.docx")

def doc_auftragsbestaetigung():
    doc = new_doc()
    eyebrow(doc, "Auftragsbestätigung")
    h1(doc, "Auftragsbestätigung [AB-2026-001]")
    address_block(doc)
    info_table(doc, [
        ("Auftragsnummer", "[AB-2026-001]"),
        ("Datum", "[TT.MM.JJJJ]"),
        ("Bezug", "Angebot [A-2026-001]"),
        ("Projektstart", "[TT.MM.JJJJ]"),
    ])
    body(doc, "Sehr geehrte:r [Name],", space_after=6)
    body(doc, "vielen Dank für Ihren Auftrag. Hiermit bestätigen wir die "
              "Beauftragung der folgenden Leistungen und freuen uns auf die "
              "Zusammenarbeit.")
    h2(doc, "Beauftragter Leistungsumfang")
    items_table(doc,
        ["Pos.", "Leistung", "Zeitraum", "Honorar"],
        [["1", "KI-Readiness-Check", "[KW XX–XX]", "[X.XXX,00 €]"],
         ["2", "Use-Case-Sprint", "[KW XX]", "[X.XXX,00 €]"]],
        widths=[1.2, 8.6, 3.5, 3.7])
    h2(doc, "Nächste Schritte")
    for t in ["Terminabstimmung für das Kick-off-Gespräch",
              "Bereitstellung der benötigten Unterlagen durch [Kunde]",
              "Start gemäß vereinbartem Zeitplan"]:
        p = para(doc, space_after=3)
        run(p, "☐  ", size=11, color=TEAL); run(p, t, size=10.5, color=BODY)
    signature(doc)
    placeholder_note(doc)
    save(doc, GA, "Auftragsbestätigung.docx")

def doc_mahnung():
    doc = new_doc()
    eyebrow(doc, "Zahlungserinnerung")
    h1(doc, "Zahlungserinnerung")
    address_block(doc)
    info_table(doc, [
        ("Rechnungsnummer", "[R-2026-001]"),
        ("Rechnungsdatum", "[TT.MM.JJJJ]"),
        ("Fällig seit", "[TT.MM.JJJJ]"),
        ("Offener Betrag", "[X.XXX,00 €]"),
    ])
    body(doc, "Sehr geehrte:r [Name],", space_after=6)
    body(doc, "vermutlich ist es Ihrer Aufmerksamkeit entgangen — die oben "
              "genannte Rechnung ist noch offen. Wir bitten Sie, den Betrag bis "
              "zum [TT.MM.JJJJ] auf unten stehendes Konto zu überweisen.")
    body(doc, "Sollten sich Ihre Zahlung und dieses Schreiben überschnitten "
              "haben, betrachten Sie die Erinnerung bitte als gegenstandslos.")
    info_table(doc, [("Bank", CO["bank"]), ("IBAN", CO["iban"]),
                     ("Verwendungszweck", "[R-2026-001]")])
    signature(doc)
    placeholder_note(doc)
    save(doc, GA, "Mahnung.docx")

def doc_signatur():
    doc = new_doc(letterhead=False)
    # eigene schlanke Ränder
    for sec in doc.sections:
        sec.top_margin = Cm(2.0)
    eyebrow(doc, "E-Mail-Signatur")
    h1(doc, "Signatur-Vorlage")
    body(doc, "Einheitliche E-Mail-Signatur für alle WirkVektor-Postfächer. "
              "Unten als formatierter Block sowie als reine Textvariante.", color=SLATE, size=10)
    # Formatierter Block
    h2(doc, "Variante A — formatiert")
    tbl = doc.add_table(rows=1, cols=1); _no_table_borders(tbl)
    cell = tbl.rows[0].cells[0]
    p0 = cell.paragraphs[0]; _p_border(p0, color=TEALHEX, sz=18, edge='left', space=8)
    p0.paragraph_format.left_indent = Cm(0.3)
    run(p0, CO["owner"], size=12, color=NAVY, bold=True, font=DISPLAY)
    for txt, col, sz, bold in [
        (CO["role"] + " · WirkVektor", SLATE, 9.5, False),
        ("", BODY, 4, False),
        (CO["claim"], TEAL, 9.5, True),
        ("", BODY, 4, False),
        ("T  " + CO["phone"] + "   ·   E  " + CO["email"], BODY, 9.5, False),
        ("W  " + CO["web"] + "   ·   " + CO["linkedin"], BODY, 9.5, False)]:
        pp = cell.add_paragraph(); pp.paragraph_format.left_indent = Cm(0.3)
        pp.paragraph_format.space_after = Pt(0)
        if txt: run(pp, txt, size=sz, color=col, bold=bold)
    para(doc, space_after=6)
    # Textvariante
    h2(doc, "Variante B — reiner Text")
    mono = para(doc, line=1.3)
    run(mono, "\n".join([
        CO["owner"], CO["role"] + " | WirkVektor",
        CO["claim"], "",
        "Tel: " + CO["phone"], "E-Mail: " + CO["email"],
        "Web: " + CO["web"], "LinkedIn: " + CO["linkedin"]]),
        size=10, color=BODY)
    placeholder_note(doc)
    save(doc, GA, "E-Mail-Signatur.docx")

def doc_visitenkarte():
    doc = new_doc(letterhead=False)
    for sec in doc.sections:
        sec.top_margin = Cm(2.0)
    eyebrow(doc, "Visitenkarte")
    h1(doc, "Visitenkarte 85 × 55 mm")
    body(doc, "Vorder- und Rückseite im Standardformat. Zum Druck an "
              "Dienstleister: 3 mm Beschnitt ergänzen.", color=SLATE, size=10)
    # Vorderseite (Navy)
    h2(doc, "Vorderseite")
    t1 = doc.add_table(rows=1, cols=1); _no_table_borders(t1)
    t1.columns[0].width = Cm(8.5)
    c = t1.rows[0].cells[0]; c.width = Cm(8.5); _set_cell_bg(c, NAVYHEX)
    c.height = Mm(55)
    _set_cell_margins(c, top=240, bottom=240, left=300, right=300)
    cp = c.paragraphs[0]; cp.paragraph_format.space_after = Pt(0)
    run(cp, "WirkVektor", size=18, color=WHITE, bold=True, font=DISPLAY)
    cp2 = c.add_paragraph(); run(cp2, CO["claim"], size=8.5, color=RGBColor.from_string(B.IMPACT_CYAN), bold=True)
    for _ in range(2): c.add_paragraph()
    n = c.add_paragraph(); n.paragraph_format.space_after = Pt(0)
    run(n, CO["owner"], size=12, color=WHITE, bold=True, font=DISPLAY)
    rr = c.add_paragraph(); run(rr, CO["role"], size=9, color=RGBColor.from_string("C7D0DE"))
    para(doc, space_after=6)
    # Rückseite (Off-White)
    h2(doc, "Rückseite")
    t2 = doc.add_table(rows=1, cols=1); _no_table_borders(t2)
    t2.columns[0].width = Cm(8.5)
    c2 = t2.rows[0].cells[0]; c2.width = Cm(8.5); _set_cell_bg(c2, B.OFF_WHITE)
    _set_cell_margins(c2, top=240, bottom=240, left=300, right=300)
    for txt, col, sz, bold in [
        ("T  " + CO["phone"], BODY, 10, False),
        ("E  " + CO["email"], BODY, 10, False),
        ("W  " + CO["web"], BODY, 10, False),
        ("   " + CO["linkedin"], SLATE, 9, False)]:
        pp = c2.add_paragraph() if txt != "T  " + CO["phone"] else c2.paragraphs[0]
        pp.paragraph_format.space_after = Pt(2)
        run(pp, txt, size=sz, color=col, bold=bold)
    placeholder_note(doc)
    save(doc, GA, "Visitenkarte.docx")

# ===========================================================================
# KATEGORIE 2 — PROJEKT & BERATUNG
# ===========================================================================
PB = "Projekt und Beratung"

def doc_protokoll():
    doc = new_doc()
    eyebrow(doc, "Protokoll")
    h1(doc, "Meeting-Protokoll")
    info_table(doc, [
        ("Anlass / Thema", "[Titel des Termins]"),
        ("Datum / Uhrzeit", "[TT.MM.JJJJ, HH:MM–HH:MM]"),
        ("Ort", "[Vor Ort / Video]"),
        ("Teilnehmende", "[Namen]"),
        ("Protokoll", CO["owner"]),
    ])
    h2(doc, "Agenda")
    for t in ["[Punkt 1]", "[Punkt 2]", "[Punkt 3]"]:
        p = para(doc, space_after=2); run(p, "•  " + t, size=10.5, color=BODY)
    h2(doc, "Besprochene Punkte & Entscheidungen")
    body(doc, "[Zusammenfassung der Diskussion und getroffener Entscheidungen.]")
    h2(doc, "Aufgaben (To-dos)")
    items_table(doc,
        ["#", "Aufgabe", "Verantwortlich", "Frist"],
        [["1", "[Aufgabe]", "[Name]", "[TT.MM.]"],
         ["2", "[Aufgabe]", "[Name]", "[TT.MM.]"],
         ["3", "[Aufgabe]", "[Name]", "[TT.MM.]"]],
        widths=[1.0, 8.8, 4.0, 3.2])
    h2(doc, "Nächster Termin")
    body(doc, "[TT.MM.JJJJ, HH:MM]", color=SLATE)
    placeholder_note(doc)
    save(doc, PB, "Meeting-Protokoll.docx")

def doc_status():
    doc = new_doc()
    eyebrow(doc, "Statusbericht")
    h1(doc, "Projektstatusbericht")
    info_table(doc, [
        ("Projekt", "[Projektname]"),
        ("Berichtszeitraum", "[KW XX–XX]"),
        ("Projektleitung", CO["owner"]),
        ("Stand", "[TT.MM.JJJJ]"),
    ])
    h2(doc, "Gesamtstatus")
    st = doc.add_table(rows=1, cols=3); _no_table_borders(st)
    for i, (lab, col) in enumerate([("Termin", TEALHEX), ("Budget", TEALHEX),
                                    ("Inhalt", "F59E0B")]):
        cell = st.rows[0].cells[i]; cell.width = Cm(5.6)
        inner = cell.add_table(rows=1, cols=1); _no_table_borders(inner)
        status_pill(inner.rows[0].cells[0], lab + ": [grün]", col)
    para(doc, space_after=4)
    h2(doc, "Erreichte Meilensteine")
    items_table(doc,
        ["Meilenstein", "Geplant", "Status"],
        [["[M1 — Readiness abgeschlossen]", "[KW XX]", "erledigt"],
         ["[M2 — Use Cases priorisiert]", "[KW XX]", "in Arbeit"],
         ["[M3 — Pilot gestartet]", "[KW XX]", "offen"]],
        widths=[9.5, 3.5, 4.0])
    h2(doc, "Risiken & Maßnahmen")
    body(doc, "[Beschreibung relevanter Risiken und geplanter Gegenmaßnahmen.]")
    h2(doc, "Nächste Schritte")
    body(doc, "[Geplante Aktivitäten im kommenden Zeitraum.]")
    placeholder_note(doc)
    save(doc, PB, "Projektstatusbericht.docx")

def doc_workshop():
    doc = new_doc()
    eyebrow(doc, "Workshop")
    h1(doc, "Workshop-Dokumentation")
    info_table(doc, [
        ("Workshop", "[Titel, z. B. KI-Use-Case-Sprint]"),
        ("Datum", "[TT.MM.JJJJ]"),
        ("Moderation", CO["owner"]),
        ("Teilnehmende", "[Namen / Rollen]"),
    ])
    h2(doc, "Ziel des Workshops")
    body(doc, "[Was sollte am Ende des Workshops erreicht sein?]")
    h2(doc, "Ablauf")
    items_table(doc,
        ["Zeit", "Programmpunkt", "Methode"],
        [["[09:00]", "[Einstieg & Ziele]", "[Plenum]"],
         ["[09:30]", "[Use-Case-Sammlung]", "[Brainstorming]"],
         ["[11:00]", "[Bewertung & Priorisierung]", "[Matrix]"],
         ["[13:00]", "[Maßnahmen & Abschluss]", "[Plenum]"]],
        widths=[2.6, 9.0, 5.4])
    h2(doc, "Ergebnisse")
    body(doc, "[Wichtigste Erkenntnisse und Festlegungen.]")
    h2(doc, "Vereinbarte Maßnahmen")
    items_table(doc,
        ["#", "Maßnahme", "Verantwortlich", "Frist"],
        [["1", "[Maßnahme]", "[Name]", "[TT.MM.]"],
         ["2", "[Maßnahme]", "[Name]", "[TT.MM.]"]],
        widths=[1.0, 8.8, 4.0, 3.2])
    placeholder_note(doc)
    save(doc, PB, "Workshop-Dokumentation.docx")

def doc_readiness():
    doc = new_doc()
    eyebrow(doc, "Ergebnisbericht")
    h1(doc, "KI-Readiness-Bericht")
    info_table(doc, [
        ("Unternehmen", "[Kundenname]"),
        ("Erstellt für", "[Ansprechpartner:in]"),
        ("Erstellt von", CO["owner"] + ", WirkVektor"),
        ("Datum", "[TT.MM.JJJJ]"),
    ])
    h2(doc, "Management Summary")
    body(doc, "[Zwei bis vier Sätze: Wo steht das Unternehmen bei der "
              "KI-Einführung, was sind die wichtigsten Hebel, welche Empfehlung "
              "geben wir?]")
    h2(doc, "Reifegrad nach Dimensionen")
    items_table(doc,
        ["Dimension", "Reifegrad", "Bewertung"],
        [["Strategie & Ziele", "[2 / 5]", "[kurz]"],
         ["Daten & Infrastruktur", "[3 / 5]", "[kurz]"],
         ["Prozesse & Use Cases", "[2 / 5]", "[kurz]"],
         ["Governance & Sicherheit", "[1 / 5]", "[kurz]"],
         ["Kompetenz & Kultur", "[3 / 5]", "[kurz]"]],
        widths=[6.5, 3.5, 7.0])
    h2(doc, "Zentrale Erkenntnisse")
    for t in ["[Erkenntnis 1]", "[Erkenntnis 2]", "[Erkenntnis 3]"]:
        p = para(doc, space_after=3)
        run(p, "▪  ", size=10.5, color=TEAL); run(p, t, size=10.5, color=BODY)
    h2(doc, "Empfehlungen & nächste Schritte")
    items_table(doc,
        ["Priorität", "Empfehlung", "Zeithorizont"],
        [["1", "[Empfehlung]", "[0–3 Monate]"],
         ["2", "[Empfehlung]", "[3–6 Monate]"],
         ["3", "[Empfehlung]", "[6–12 Monate]"]],
        widths=[2.5, 10.5, 4.0])
    placeholder_note(doc)
    save(doc, PB, "KI-Readiness-Bericht.docx")

def doc_usecase():
    doc = new_doc()
    eyebrow(doc, "Use-Case-Steckbrief")
    h1(doc, "[Name des Use Case]")
    info_table(doc, [
        ("Use Case", "[Kurzname]"),
        ("Fachbereich", "[Bereich]"),
        ("Erstellt am", "[TT.MM.JJJJ]"),
        ("Priorität", "[Hoch / Mittel / Gering]"),
    ])
    h2(doc, "Ziel & Beschreibung")
    body(doc, "[Welches Problem wird gelöst? Wie sieht die KI-Unterstützung "
              "konkret aus?]")
    h2(doc, "Bewertung")
    items_table(doc,
        ["Kriterium", "Einschätzung", "Kommentar"],
        [["Erwarteter Nutzen", "[Hoch]", "[kurz]"],
         ["Umsetzungsaufwand", "[Mittel]", "[kurz]"],
         ["Risiko / Compliance", "[Gering]", "[kurz]"],
         ["Datenverfügbarkeit", "[Gut]", "[kurz]"]],
        widths=[5.5, 4.0, 7.5])
    h2(doc, "Datenbedarf & Voraussetzungen")
    body(doc, "[Welche Daten, Systeme und Freigaben werden benötigt?]")
    h2(doc, "Nächster Schritt")
    body(doc, "[z. B. Proof of Concept, Datenfreigabe, Pilotierung]", color=SLATE)
    placeholder_note(doc)
    save(doc, PB, "Use-Case-Steckbrief.docx")

# ===========================================================================
# KATEGORIE 3 — INTERN & MARKETING
# ===========================================================================
IM = "Intern und Marketing"

def doc_onboarding():
    doc = new_doc()
    eyebrow(doc, "Onboarding")
    h1(doc, "Onboarding-Checkliste")
    info_table(doc, [
        ("Neue:r Mitarbeiter:in", "[Name]"),
        ("Rolle", "[Position]"),
        ("Startdatum", "[TT.MM.JJJJ]"),
        ("Pate / Buddy", "[Name]"),
    ])
    groups = [
        ("Vor dem ersten Tag", ["Arbeitsvertrag unterschrieben", "Hardware bestellt",
            "Zugänge & Konten vorbereitet", "Willkommens-E-Mail versendet"]),
        ("Tag 1", ["Begrüßung & Rundgang", "Arbeitsplatz & Technik einrichten",
            "Team vorstellen", "Zugang zu Vault & Tools"]),
        ("Woche 1", ["WirkVektor-Konzept & Schreibstil lesen", "Designsystem kennenlernen",
            "Erste Aufgabe übernehmen", "Feedback-Gespräch terminieren"]),
        ("Erster Monat", ["Kundenprozesse verstehen", "An Erstgespräch teilnehmen",
            "30-Tage-Feedback", "Weiterbildungsplan abstimmen"]),
    ]
    for title, items in groups:
        h2(doc, title)
        for it in items:
            p = para(doc, space_after=3)
            run(p, "☐  ", size=12, color=TEAL); run(p, it, size=10.5, color=BODY)
    placeholder_note(doc)
    save(doc, IM, "Onboarding-Checkliste.docx")

def doc_onepager():
    doc = new_doc()
    eyebrow(doc, "One-Pager")
    p = para(doc, space_after=2); run(p, "Die Architektur wirksamer KI.", size=24, color=NAVY, bold=True, font=DISPLAY)
    body(doc, CO["claim"], color=TEAL, size=11, bold=True)
    h2(doc, "Das Problem")
    body(doc, "Viele mittelständische Unternehmen probieren KI aus, steuern sie "
              "aber nicht. Es fehlt an priorisierten Use Cases, an Governance und "
              "an messbarem Nutzen.")
    h2(doc, "Unser Ansatz")
    body(doc, "WirkVektor bringt KI kontrolliert in die Unternehmenspraxis — mit "
              "klarer Wirkung, sauberer Governance und messbarem Nutzen. "
              "Schwerpunkt: Informationssicherheit und Datenschutz als Kernkompetenz.")
    h2(doc, "Leistungspakete")
    cols = doc.add_table(rows=1, cols=len(B.PAKETE)); _no_table_borders(cols)
    for i, name in enumerate(B.PAKETE):
        cell = cols.rows[0].cells[i]; cell.width = Cm(3.4)
        _set_cell_bg(cell, B.OFF_WHITE); _set_cell_margins(cell)
        pp = cell.paragraphs[0]; pp.paragraph_format.space_after = Pt(0)
        run(pp, f"0{i+1}", size=11, color=TEAL, bold=True, font=DISPLAY)
        pn = cell.add_paragraph(); pn.paragraph_format.space_after = Pt(0)
        run(pn, name, size=9, color=NAVY, bold=True)
    para(doc, space_after=6)
    h2(doc, "Methodik in 6 Phasen")
    body(doc, "  →  ".join(m[1] for m in B.METHODIK), color=SLATE, size=10)
    h2(doc, "Kontakt")
    body(doc, f"{CO['owner']} · {CO['role']} · {CO['email']} · {CO['web']}",
         color=BODY, size=10)
    placeholder_note(doc)
    save(doc, IM, "One-Pager.docx")

def doc_casestudy():
    doc = new_doc()
    eyebrow(doc, "Case Study")
    h1(doc, "[Kundenname]: [Ergebnis in einem Satz]")
    info_table(doc, [
        ("Branche", "[Branche]"),
        ("Unternehmensgröße", "[X Mitarbeitende]"),
        ("Leistung", "[z. B. KI-Readiness-Check + Pilot]"),
        ("Zeitraum", "[X Wochen]"),
    ])
    h2(doc, "Ausgangslage")
    body(doc, "[Welche Herausforderung hatte der Kunde?]")
    h2(doc, "Vorgehen")
    body(doc, "[Wie sind wir vorgegangen — in zwei bis drei Sätzen?]")
    h2(doc, "Lösung")
    body(doc, "[Was wurde umgesetzt?]")
    h2(doc, "Ergebnis")
    res = doc.add_table(rows=1, cols=3); _no_table_borders(res)
    for i, (val, lab) in enumerate([("[–40 %]", "[Bearbeitungszeit]"),
                                    ("[3×]", "[schnellere Auswahl]"),
                                    ("[100 %]", "[EU-AI-Act-konform]")]):
        cell = res.rows[0].cells[i]; cell.width = Cm(5.6); _set_cell_margins(cell)
        pv = cell.paragraphs[0]; run(pv, val, size=22, color=TEAL, bold=True, font=DISPLAY)
        pl = cell.add_paragraph(); pl.paragraph_format.space_after = Pt(0)
        run(pl, lab, size=9, color=SLATE)
    para(doc, space_after=6)
    h2(doc, "Kundenstimme")
    q = para(doc); q.paragraph_format.left_indent = Cm(0.3)
    _p_border(q, color=TEALHEX, sz=18, edge='left', space=8)
    run(q, "„[Zitat des Kunden.]“", size=12, color=NAVY, italic=True, font=DISPLAY)
    qa = para(doc); run(qa, "— [Name, Funktion]", size=9.5, color=SLATE)
    placeholder_note(doc)
    save(doc, IM, "Case-Study.docx")

def doc_presse():
    doc = new_doc()
    eyebrow(doc, "Pressemitteilung")
    body(doc, "[Ort], [TT.MM.JJJJ]", color=SLATE, size=9.5)
    p = para(doc, space_after=4); run(p, "[Aussagekräftige Schlagzeile in einer Zeile]", size=20, color=NAVY, bold=True, font=DISPLAY)
    p2 = para(doc, space_after=10); run(p2, "[Unterzeile, die die Schlagzeile konkretisiert.]", size=12, color=SLATE, font=DISPLAY)
    body(doc, "[Einleitung: Die wichtigste Information zuerst — wer, was, wann, "
              "wo, warum. Ein bis zwei Sätze.]", bold=True)
    body(doc, "[Hauptteil: Hintergrund, Kontext und Details. Sachlich und "
              "belegbar, ohne werbliche Übertreibung.]")
    body(doc, "[Optionales Zitat:] „[Aussage von Sebastian Schucht].“, sagt "
              + CO["owner"] + ", " + CO["role"] + " von WirkVektor.")
    body(doc, "[Abschließender Absatz mit Ausblick oder Einordnung.]")
    h2(doc, "Über WirkVektor")
    body(doc, "WirkVektor ist eine Beratungsgesellschaft für die sichere, "
              "strukturierte und wirksame Einführung von Künstlicher Intelligenz "
              "in mittelständischen Unternehmen. Schwerpunkte sind KI-Strategie, "
              "Governance und EU AI Act sowie die produktive Einführung "
              "generativer KI.", size=9.5, color=SLATE)
    h2(doc, "Pressekontakt")
    body(doc, f"{CO['owner']} · {CO['email']} · {CO['phone']} · {CO['web']}",
         size=9.5, color=BODY)
    placeholder_note(doc)
    save(doc, IM, "Pressemitteilung.docx")

# ===========================================================================
ALL = [
    doc_briefbogen, doc_angebot, doc_rechnung, doc_auftragsbestaetigung,
    doc_mahnung, doc_signatur, doc_visitenkarte,
    doc_protokoll, doc_status, doc_workshop, doc_readiness, doc_usecase,
    doc_onboarding, doc_onepager, doc_casestudy, doc_presse,
]

def build():
    print("Erzeuge Word-Dokumente:")
    for fn in ALL:
        fn()
    print(f"Fertig — {len(ALL)} Dokumente.")

if __name__ == "__main__":
    build()
